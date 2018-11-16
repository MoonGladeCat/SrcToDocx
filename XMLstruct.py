#!/usr/bin/python3
# -*- coding:utf-8 -*-
'''
Created on Nov 14, 2018

@author: kimi
'''

import xml.etree.ElementTree as ET
import os
from docx import Document

defaultString = 'None'

class MyMember(object):
    def __init__(self, memberElement):
        self.type       = ''.join(memberElement.find('./type').itertext())
        self.definition = memberElement.findtext('./definition', defaultString)
        self.argsString = memberElement.findtext('./argsstring', defaultString)
        self.name       = memberElement.findtext('./name', defaultString)
        self.brief      = memberElement.findtext('./briefdescription', defaultString)
        self.brief = self.brief.strip()
        if len(self.brief) == 0:
            self.brief = defaultString
        self.detail     = ''.join(memberElement.find('./detaileddescription').itertext())
        self.detail     = self.detail.strip()
        if len(self.detail) == 0:
            self.detail = defaultString        
        
    def __str__(self):
        string = 'Type: {0}  Name: {1}  Definiton: {2}'.format(self.type, self.name, self.definition)
        return string

class MyStruct(object):
    def __init__(self, structElement):
        self.id     = structElement.get('id', defaultString)
        self.kind   = structElement.get('kind', defaultString)
        self.name   = structElement.findtext('./compoundname', defaultString)
        self.brief  = structElement.findtext('./briefdescription', defaultString)
        self.brief  = self.brief.strip()
        if len(self.brief) == 0:
            self.brief = defaultString
        self.detail = ''.join(structElement.find('./detaileddescription').itertext())
        self.detail = self.detail.strip()
        if len(self.detail) == 0:
            self.detail = defaultString  
        self.members = [];
        for memberEle in structElement.findall('./sectiondef/memberdef'):
            member = MyMember(memberEle)
            if member.name !=  defaultString and member not in self.members:
                self.members.append(member)
  
    def __str__(self):
        string = 'Id: {0} Kind: {1} Name: {2}'.format(self.id, self.kind, self.name)
        string = string + os.linesep + repr(self.brief) + os.linesep + repr(self.detail)
        for member in self.members:
            string = string + os.linesep + member.__str__()
        return string
    
class XMLStruct(object):
    def __init__(self, xmlFile):
        self.path = os.path.dirname(os.path.abspath(xmlFile))
        self.name = os.path.splitext(os.path.basename(xmlFile))[0]
        self.structs = []           
        root = ET.parse(xmlFile).getroot()        
        for structEle in root.findall('./compounddef'):
            struct = MyStruct(structEle)
            if struct.name != defaultString and struct not in self.structs:
                self.structs.append(struct)
                print(struct)
           
    def genDocx(self, docxPath = None):
        if docxPath is None:
            docxPath = self.path + os.path.sep + self.name + '.docx'
          
        document = Document()
        
        for struct in self.structs:
            content = struct.name + ' Struct Reference'            
            document.add_paragraph(content, style = 'Heading 2')
            
            document.add_paragraph('', style = 'Body Text')
            for member in struct.members:
                content = member.type + ' ' + member.name + member.argsString
                document.add_paragraph(content, style = 'List Bullet')
            document.add_paragraph('', style = 'Body Text')    
            
            content = 'Detailed Description'
            document.add_paragraph(content, style = 'Heading 2')
            document.add_paragraph('', style = 'Body Text')
            document.add_paragraph(struct.detail, style = 'Body Text')
            document.add_paragraph('', style = 'Body Text')
            
            
            content = 'Filed Documentation'
            document.add_paragraph(content, style = 'Heading 2')
            document.add_paragraph('', style = 'Body Text')
            for member in struct.members:
                content = member.type + ' ' + member.name + member.argsString
                content = content + ' : ' + member.detail
                document.add_paragraph(content, style = 'List Bullet')
                
            
            document.add_page_break()
        document.save(docxPath)
        
                
if __name__ == '__main__':
    
    xmlFile  = './src/xml/struct__UM__T__SCHEDULE__.xml'

    xmlStruct =  XMLStruct(xmlFile)
    
    xmlStruct.genDocx()
    
        
        